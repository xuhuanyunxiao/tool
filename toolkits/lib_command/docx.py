# -*- coding: utf-8 -*-
"""
Created on Thu Jan 25 18:31:52 2018

@author: Administrator
"""
#%%
from docx import Document
from docx.shared import  Pt
from docx.oxml.ns import  qn
from docx.shared import Inches
from docx.shared import RGBColor

#%% 打开文档
document = Document()

#%% 加入不同等级的标题
document.add_heading('Document Title',0)
document.add_heading(u'二级标题',1)
document.add_heading(u'二级标题',2)

#%% 添加文本
paragraph = document.add_paragraph(u'添加了文本')
# 增加两个段落
file.add_paragraph('我喜欢你')
file.add_paragraph('的眼')
# 在段落中新增文字
p = file.add_paragraph()
p.add_run('我喜欢你')
p.add_run('的眼')

#设置字号
run = paragraph.add_run(u'设置字号')
run.font.size=Pt(24)

#设置字体
run = paragraph.add_run('Set Font,')
run.font.name='Consolas'

#设置中文字体
run = paragraph.add_run(u'设置中文字体，')
# 通过file. add_paragraph()生成的段落是无法直接设置字体的，所以不能直接写font = p.font
run.font.name=u'宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
# file.styles['Normal'].font.name = u'微软雅黑'

font = run.font

# 设置字体颜色
font.color.rgb = RGBColor(54, 95,145)

#设置斜体
run = paragraph.add_run(u'斜体、')
run.italic = True

#设置粗体
run = paragraph.add_run(u'粗体').bold = True

# 设置下划线
font.underline = True

#增加引用
document.add_paragraph('Intense quote', style='Intense Quote')

#增加有序列表
document.add_paragraph(
    u'有序列表元素1',style='List Number'
)
document.add_paragraph(
    u'有序列别元素2',style='List Number'
)

#增加无序列表
document.add_paragraph(
    u'无序列表元素1',style='List Bullet'
)
document.add_paragraph(
    u'无序列表元素2',style='List Bullet'
)

#增加图片（此处使用相对位置）
document.add_picture('jdb.jpg',width=Inches(1.25))

#增加表格
# 表格样式
table = document.add_table(rows=3,cols=3, style = '')
hdr_cells=table.rows[0].cells
hdr_cells[0].text="第一列"
hdr_cells[1].text="第二列"
hdr_cells[2].text="第三列"

hdr_cells = table.rows[1].cells
hdr_cells[0].text = '2'
hdr_cells[1].text = 'aerszvfdgx'
hdr_cells[2].text = 'abdzfgxfdf'

hdr_cells = table.rows[2].cells
hdr_cells[0].text = '3'
hdr_cells[1].text = 'cafdwvaef'
hdr_cells[2].text = 'aabs zfgf'

#增加分页
document.add_page_break()

#增加图片
document.add_picture('pic.png', width = Inches(3), height = Inches(3))
# 图片默认放在最左边，如果需要居中放置
last_paragraph = document.paragraph[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


#保存文件
document.save('demo.docx')

for paragraph in document.paragraph:
    print(paragraph) # 该Word文档中的所有段落
#%%

p = file.add_paragraph()
# 通过file. add_paragraph()生成的段落是无法直接设置字体的，所以不能直接写font = p.font
run = p.add_run('我喜欢你')
font = run.font

# 字体大小，单位为“磅”，磅是字体大小的单位point的音译，缩写为pt
font.size=Pt(24) 
# 设置字体颜色
font.color.rgb = RGBColor(54, 95,145)
#设置斜体
font.italic = True
#设置粗体
font.bold = True
# 设置下划线
font.underline = True


# 对齐方式
p = file.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # 居中对齐
p.add_run('而你在想我')

# 首行缩进, Inches括号里的数字为缩进量，单位为英寸
p.paragraph_format.first_line_indent = Inches(0.32)
# 小四号对应0.16英寸，而缩进一般是2个字符长度，
# 所以设置0.32英寸的首行缩进量，还可以对其进行适当微调。
# 对于其他大小的字体，可以参照此方法计算首行缩进量。

# 行距和段落间距
p.paragraph_format.line_spacing = Pt(16)
# 行距，16磅对应三号字体大小，14磅对应四号字体
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(14)

# 项目符号和编号
# 在添加段落时设置参数style=’序号格式’，
# “List Bullet”表示项目符号，“List Number”表示数字编号
file.add_paragraph('项目符号示例文字', style = 'List Bullet')
file.add_paragraph('数字编号示例文字', style = 'List Number')